"""Document processing pipeline — parse, chunk, and index documents."""

import re
import uuid
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.logging import get_logger
from app.models.database import Document
from app.services.vector_store import vector_store

logger = get_logger(__name__)


async def process_document(document_id: str, file_path: str, db: AsyncSession):
    """Parse, chunk, and index a document. Updates the DB record on completion or failure.

    Called as a background task with its own DB session (passed from the route).
    """
    path = Path(file_path)
    try:
        logger.info("processing_started", document_id=document_id, file_path=file_path)

        # 1. Parse document to raw text
        text = parse_document(path)
        if not text.strip():
            raise ValueError("Document is empty or could not be parsed")

        # 2. Chunk the text
        chunks = chunk_text(
            text,
            chunk_size=settings.chunk_size,
            overlap=settings.chunk_overlap,
        )
        if not chunks:
            raise ValueError("No chunks produced from document")

        # 3. Get the original filename from DB
        result = await db.execute(select(Document).where(Document.id == document_id))
        doc = result.scalar_one_or_none()
        if not doc:
            raise ValueError(f"Document {document_id} not found in database")

        # 4. Add chunks to vector store
        count = vector_store.add_chunks(document_id, doc.filename, chunks)

        # 5. Update DB record
        doc.status = "ready"
        doc.chunk_count = count
        await db.commit()

        logger.info(
            "processing_complete",
            document_id=document_id,
            chunk_count=count,
        )

    except Exception as e:
        logger.error("processing_failed", document_id=document_id, error=str(e))
        try:
            result = await db.execute(select(Document).where(Document.id == document_id))
            doc = result.scalar_one_or_none()
            if doc:
                doc.status = "failed"
                doc.error_message = str(e)
                await db.commit()
        except Exception as db_err:
            logger.error("status_update_failed", document_id=document_id, error=str(db_err))


def parse_document(file_path: Path) -> str:
    """Extract text from a document based on its file extension."""
    ext = file_path.suffix.lower()

    if ext in (".md", ".txt"):
        return file_path.read_text(encoding="utf-8")

    elif ext == ".pdf":
        return _parse_pdf(file_path)

    elif ext == ".docx":
        return _parse_docx(file_path)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(file_path: Path) -> str:
    """Extract text from a PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {i}]\n{text.strip()}")
    return "\n\n".join(pages)


def _parse_docx(file_path: Path) -> str:
    """Extract text from a DOCX using python-docx."""
    import docx

    doc = docx.Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def chunk_text(
    text: str,
    chunk_size: int = 500,
    overlap: int = 50,
) -> list[dict]:
    """Split text into overlapping chunks, respecting paragraph and sentence boundaries.

    Uses character-based sizing (1 token ~ 4 chars) for simplicity.
    Returns list of dicts: {text, chunk_index, page_or_section}.
    """
    char_limit = chunk_size * 4
    char_overlap = overlap * 4

    # Split into paragraphs first
    paragraphs = re.split(r"\n\n+", text.strip())

    # Build chunks by accumulating paragraphs
    chunks: list[dict] = []
    current_text = ""
    current_section = None
    chunk_index = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Detect section headings (markdown # or [Page N] markers)
        heading_match = re.match(r"^(#{1,6})\s+(.+)", para)
        page_match = re.match(r"^\[Page\s+(\d+)\]", para)
        if heading_match:
            current_section = heading_match.group(2).strip()
        elif page_match:
            current_section = f"Page {page_match.group(1)}"

        # If adding this paragraph would exceed the limit, flush current chunk
        if current_text and len(current_text) + len(para) + 2 > char_limit:
            chunks.append(
                {
                    "text": current_text.strip(),
                    "chunk_index": chunk_index,
                    "page_or_section": current_section,
                }
            )
            chunk_index += 1

            # Carry overlap from the end of the current chunk
            if char_overlap > 0 and len(current_text) > char_overlap:
                current_text = current_text[-char_overlap:]
            else:
                current_text = ""

        # If a single paragraph exceeds the limit, split it further
        if len(para) > char_limit:
            # Flush any accumulated text first
            if current_text.strip():
                chunks.append(
                    {
                        "text": current_text.strip(),
                        "chunk_index": chunk_index,
                        "page_or_section": current_section,
                    }
                )
                chunk_index += 1
                current_text = ""

            sub_chunks = _split_long_paragraph(para, char_limit, char_overlap)
            for sc in sub_chunks:
                chunks.append(
                    {
                        "text": sc.strip(),
                        "chunk_index": chunk_index,
                        "page_or_section": current_section,
                    }
                )
                chunk_index += 1

            # Carry overlap from the last sub-chunk
            last = sub_chunks[-1] if sub_chunks else ""
            if char_overlap > 0 and len(last) > char_overlap:
                current_text = last[-char_overlap:]
            else:
                current_text = ""
        else:
            if current_text:
                current_text += "\n\n" + para
            else:
                current_text = para

    # Flush remaining text
    if current_text.strip():
        chunks.append(
            {
                "text": current_text.strip(),
                "chunk_index": chunk_index,
                "page_or_section": current_section,
            }
        )

    return chunks


def _split_long_paragraph(text: str, char_limit: int, char_overlap: int) -> list[str]:
    """Split a paragraph that exceeds char_limit using sentence boundaries, then newlines, then hard split."""
    # Try splitting on sentences first
    sentences = re.split(r"(?<=[.!?])\s+", text)
    if len(sentences) > 1:
        return _accumulate_splits(sentences, char_limit, char_overlap)

    # Try splitting on newlines
    lines = text.split("\n")
    if len(lines) > 1:
        return _accumulate_splits(lines, char_limit, char_overlap)

    # Hard split at char_limit
    result = []
    start = 0
    while start < len(text):
        end = min(start + char_limit, len(text))
        result.append(text[start:end])
        start = end - char_overlap if char_overlap > 0 else end
    return result


def _accumulate_splits(
    parts: list[str], char_limit: int, char_overlap: int
) -> list[str]:
    """Accumulate text parts into chunks that stay under char_limit."""
    chunks = []
    current = ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if current and len(current) + len(part) + 1 > char_limit:
            chunks.append(current)
            # Carry overlap
            if char_overlap > 0 and len(current) > char_overlap:
                current = current[-char_overlap:] + " " + part
            else:
                current = part
        else:
            current = (current + " " + part).strip() if current else part
    if current:
        chunks.append(current)
    return chunks


async def load_sample_documents():
    """Load sample documents from the sample_docs directory if not already indexed.

    Creates its own DB session since this is called from the app lifespan.
    """
    from app.models.database import async_session as session_factory

    sample_dir = settings.sample_docs_dir
    if not sample_dir.exists():
        logger.info("sample_docs_dir_missing", path=str(sample_dir))
        return

    sample_files = list(sample_dir.glob("*"))
    if not sample_files:
        logger.info("no_sample_docs_found")
        return

    async with session_factory() as db:
        # Check if any documents already exist
        result = await db.execute(select(Document))
        existing = result.scalars().all()
        if existing:
            logger.info("sample_docs_skipped", reason="documents_already_exist", count=len(existing))
            return

        logger.info("loading_sample_docs", count=len(sample_files))

        for file_path in sorted(sample_files):
            ext = file_path.suffix.lower()
            if ext not in settings.supported_extensions:
                continue

            try:
                content = file_path.read_bytes()
                document_id = str(uuid.uuid4())
                doc = Document(
                    id=document_id,
                    filename=file_path.name,
                    file_size=len(content),
                    status="processing",
                )
                db.add(doc)
                await db.commit()
                await db.refresh(doc)

                await process_document(document_id, str(file_path), db)

                logger.info("sample_doc_loaded", filename=file_path.name, document_id=document_id)

            except Exception as e:
                logger.error("sample_doc_failed", filename=file_path.name, error=str(e))
